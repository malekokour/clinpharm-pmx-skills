#!/usr/bin/env python3
"""Build polished DOCX artifacts from canonical ClinPharm PMx Skills Markdown.

Author: ClinPharm PMx Skills contributors
Date: 2026-07-29
Dependencies: python-docx
"""

from __future__ import annotations

import argparse
import re
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = "17324D"
TEAL = "0B7A75"
TEAL_LIGHT = "E9F4F2"
NAVY_LIGHT = "EAF0F6"
GRAY = "5F6B76"
LINE = "CFD8E3"
WHITE = "FFFFFF"
FONT = "Aptos"
MONO = "Aptos Mono"
PACKAGE_TIME = (2026, 7, 30, 12, 0, 0)


def normalize_docx_package(path: Path) -> None:
    """Rewrite ZIP metadata so identical source produces identical bytes."""
    with zipfile.ZipFile(path) as source:
        members = [(name, source.read(name)) for name in sorted(source.namelist())]
    with NamedTemporaryFile(suffix=".docx", dir=path.parent, delete=False) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(temporary, "w") as archive:
            for name, data in members:
                info = zipfile.ZipInfo(name, PACKAGE_TIME)
                # ZipInfo otherwise derives this field from the host OS, which
                # makes an equivalent package differ between Windows and Unix.
                info.create_system = 3
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o100644 << 16
                archive.writestr(info, data)
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_bottom_border(paragraph, color: str = TEAL, size: int = 16) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.06

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 10),
        ("Heading 1", 18, INK, 12, 6),
        ("Heading 2", 13, TEAL, 10, 4),
        ("Heading 3", 10.5, INK, 7, 2),
    ):
        style = document.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    document.styles["Title"].font.bold = True
    document.styles["Heading 1"].font.bold = True

    for list_name in ("List Bullet", "List Number"):
        list_style = document.styles[list_name]
        list_style.font.name = FONT
        list_style.font.size = Pt(9.25)
        list_style.font.color.rgb = RGBColor.from_string(INK)
        list_style.paragraph_format.space_after = Pt(2.5)

    if "ClinPharm PMx Skills Quote" not in document.styles:
        quote = document.styles.add_style("ClinPharm PMx Skills Quote", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = normal
        quote.font.name = FONT
        quote.font.size = Pt(9.25)
        quote.font.italic = True
        quote.font.color.rgb = RGBColor.from_string(GRAY)
        quote.paragraph_format.left_indent = Inches(0.18)
        quote.paragraph_format.right_indent = Inches(0.18)
        quote.paragraph_format.space_before = Pt(4)
        quote.paragraph_format.space_after = Pt(7)

    if "ClinPharm PMx Skills Code" not in document.styles:
        code = document.styles.add_style("ClinPharm PMx Skills Code", WD_STYLE_TYPE.PARAGRAPH)
        code.base_style = normal
        code.font.name = MONO
        code.font.size = Pt(8.3)
        code.font.color.rgb = RGBColor.from_string(INK)
        code.paragraph_format.left_indent = Inches(0.15)
        code.paragraph_format.right_indent = Inches(0.15)
        code.paragraph_format.space_after = Pt(2)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.clear()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("CLINPHARM SKILLS  /  PORTABLE WORK CONTEXT")
    run.font.name = FONT
    run.font.size = Pt(7.2)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    add_bottom_border(header_p, LINE, 6)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("ClinPharm PMx Skills  •  Public or synthetic examples only  •  ")
    run.font.name = FONT
    run.font.size = Pt(7.2)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(footer_p, "PAGE")


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = MONO
            run.font.size = Pt(8.7)
            run.font.color.rgb = RGBColor.from_string(TEAL)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            run.font.color.rgb = RGBColor.from_string(TEAL)
            run.underline = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def strip_frontmatter(lines: list[str]) -> tuple[list[str], list[str]]:
    if not lines or lines[0].strip() != "---":
        return [], lines
    metadata: list[str] = []
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            return metadata, lines[index + 1 :]
        metadata.append(lines[index])
    return [], lines


def join_list_continuations(lines: list[str]) -> list[str]:
    """Join indented Markdown list continuations into their owning list item."""
    joined: list[str] = []
    for line in lines:
        if (
            re.match(r"^\s{2,}\S", line)
            and joined
            and re.match(r"^\s*(?:[-*]|\d+\.)\s+", joined[-1])
        ):
            joined[-1] = joined[-1].rstrip() + " " + line.strip()
        else:
            joined.append(line)
    return joined


def add_metadata_card(document: Document, metadata: list[str]) -> None:
    selected = []
    for line in metadata:
        if re.match(r"^(version|updated_at|status|data_classification|review_due|project_slug):", line):
            key, value = line.split(":", 1)
            selected.append((key.replace("_", " ").title(), value.strip().strip('"')))
    if not selected:
        return
    paragraph = document.add_paragraph()
    set_paragraph_shading(paragraph, NAVY_LIGHT)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_together = True
    for index, (label, value) in enumerate(selected):
        if index and index % 3 == 0:
            paragraph.add_run().add_break()
        elif index:
            separator = paragraph.add_run("   •   ")
            separator.font.name = FONT
            separator.font.size = Pt(7.4)
            separator.font.color.rgb = RGBColor.from_string(GRAY)
        label_run = paragraph.add_run(label.upper() + " ")
        label_run.font.name = FONT
        label_run.font.size = Pt(7.2)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor.from_string(TEAL)
        value_run = paragraph.add_run(value)
        value_run.font.name = FONT
        value_run.font.size = Pt(7.8)
        value_run.font.bold = True
        value_run.font.color.rgb = RGBColor.from_string(INK)


def new_numbering_id(document: Document, base_num_id: int = 5) -> int:
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_id = max(existing, default=0) + 1
    base = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_node = OxmlElement("w:numId")
    num_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_node])


def add_table(document: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index in range(width):
            cell = row.cells[col_index]
            value = values[col_index] if col_index < len(values) else ""
            cell.text = ""
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, INK)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(8)
                if row_index == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def markdown_to_docx(source: Path, output: Path) -> None:
    raw_lines = source.read_text(encoding="utf-8").splitlines()
    metadata, lines = strip_frontmatter(raw_lines)
    lines = join_list_continuations(lines)
    document = Document()
    style_document(document)
    document.core_properties.title = source.stem.replace("-", " ")
    document.core_properties.author = "ClinPharm PMx Skills"
    document.core_properties.subject = "Portable professional and project context"
    document.core_properties.keywords = "ClinPharm PMx Skills, pharmaceutical, professional context"
    document.core_properties.comments = "Generated from canonical Markdown."
    document.core_properties.created = datetime(2026, 7, 30, tzinfo=UTC)
    document.core_properties.modified = datetime(2026, 7, 30, tzinfo=UTC)
    document.core_properties.last_modified_by = "ClinPharm PMx Skills"

    first_title = True
    paragraph_buffer: list[str] = []
    table_buffer: list[list[str]] = []
    code_buffer: list[str] = []
    in_code = False
    active_numbering_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = document.add_paragraph()
            add_inline(p, " ".join(line.strip() for line in paragraph_buffer))
            paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            useful = [
                row
                for row in table_buffer
                if not all(re.fullmatch(r":?-{3,}:?", value.strip()) for value in row)
            ]
            if useful:
                add_table(document, useful)
            table_buffer = []

    def flush_code() -> None:
        nonlocal code_buffer
        if code_buffer:
            p = document.add_paragraph()
            p.style = "ClinPharm PMx Skills Code"
            set_paragraph_shading(p, NAVY_LIGHT)
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.keep_together = True
            p.add_run("\n".join(code_buffer))
            code_buffer = []

    for raw in [*lines, ""]:
        line = raw.rstrip()
        is_numbered = bool(re.match(r"^\s*\d+\.\s+(.+)$", line))
        if not is_numbered:
            active_numbering_id = None
        if line.startswith("```"):
            flush_paragraph()
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            table_buffer.append([part.strip() for part in line.strip("|").split("|")])
            continue
        flush_table()
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if source.name == "Pharma-Work-Context.md" and text == "Final verification":
                document.add_page_break()
            if level == 1 and first_title:
                p = document.add_paragraph(style="Title")
                add_inline(p, text)
                add_bottom_border(p, TEAL, 20)
                subtitle = document.add_paragraph("Portable Agent Skills that teach AI how pharmaceutical professionals work.")
                subtitle.paragraph_format.space_after = Pt(10)
                for run in subtitle.runs:
                    run.font.name = FONT
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor.from_string(GRAY)
                add_metadata_card(document, metadata)
                first_title = False
            else:
                style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
                p = document.add_paragraph(style=style)
                add_inline(p, text)
            continue
        if re.fullmatch(r"\s*---\s*", line):
            flush_paragraph()
            p = document.add_paragraph()
            add_bottom_border(p, LINE, 6)
            p.paragraph_format.space_after = Pt(2)
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            p = document.add_paragraph(style="List Number")
            if active_numbering_id is None:
                active_numbering_id = new_numbering_id(document)
            apply_numbering(p, active_numbering_id)
            add_inline(p, numbered.group(1))
            continue
        quote = re.match(r"^\s*>\s?(.*)$", line)
        if quote:
            flush_paragraph()
            p = document.add_paragraph(style="ClinPharm PMx Skills Quote")
            add_inline(p, quote.group(1))
            continue
        if not line.strip():
            flush_paragraph()
            continue
        paragraph_buffer.append(line)

    flush_paragraph()
    flush_table()
    flush_code()

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.widow_control = True
        else:
            paragraph.paragraph_format.widow_control = True

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    normalize_docx_package(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path, help="Canonical Markdown source")
    parser.add_argument("output", type=Path, help="Generated DOCX output")
    args = parser.parse_args()
    markdown_to_docx(args.source.resolve(), args.output.resolve())
    print(f"Built {args.output}")


if __name__ == "__main__":
    main()
